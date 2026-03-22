"""DocxParser — python-docx 提取段落 + 表格，作为单页文档返回。"""

from __future__ import annotations

from pathlib import Path

from src.ingest.pdf_analyzer import PageContent, ParsedDocument


class DocxParser:
    def parse(self, file_path: Path) -> ParsedDocument:
        from docx import Document

        doc = Document(str(file_path))
        lines: list[str] = []

        for para in doc.paragraphs:
            stripped = para.text.strip()
            if stripped:
                lines.append(stripped)

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if row_text:
                    lines.append(row_text)

        text = "\n".join(lines)
        page = PageContent(page_num=1, text=text, headers=[], is_ocr=False)
        return ParsedDocument(
            file_path=file_path,
            file_name=file_path.name,
            total_pages=1,
            is_scanned=False,
            pages=[page],
        )
